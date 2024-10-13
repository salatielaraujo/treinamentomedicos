
# Instale com:
# pip install crewai
# pip install crewai-tools
# pip install python-docx
# pip install langchain
# pip install deep-translator

import streamlit as st
from crewai import Agent, Task, Crew, Process
import os
from crewai_tools import ScrapeWebsiteTool, SerperDevTool
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from docx import Document
from io import BytesIO
import base64
from io import BytesIO
from langchain import PromptTemplate
from deep_translator import GoogleTranslator 

st.set_page_config(  layout="wide" )

load_dotenv()

# Verifique se a chave foi carregada corretamente
openai_api_key = os.getenv("OPENAI_API_KEY")

if openai_api_key is None:
    raise ValueError("A chave da API OPENAI_API_KEY não foi encontrada no ambiente.")

# LLM object and API Key
os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")
os.environ["SERPER_API_KEY"] = os.getenv("SERPER_API_KEY")

def translate_text(text, source_lang='en', target_lang='pt'):
    print(f"Traduzindo: {text} ")
    translated_text = GoogleTranslator(source=source_lang, target=target_lang).translate(text)
    print(f"Texto traduzido: {translated_text} ")
    return translated_text

def generate_docx(result):

    doc = Document()
    doc.add_heading('Diagnóstico de saúde e recomendações de tratamento', 0)

    # Se result for um dicionário, convertê-lo para uma string formatada
    result = str(result)  # Converte qualquer outro tipo em string
   
    doc.add_paragraph(result.replace('\r\n', ' ').replace('\n', ' '))
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def get_download_link(bio, filename):
    b64 = base64.b64encode(bio.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Diagnosis and Treatment Plan</a>'

# Adicionar CSS customizado para centralizar o título e criar boxes ao redor das colunas
st.markdown("""
    <style>
    /* Centralizar o título */
    h1 {
        text-align: center;
    }

    /* Adicionar borda em volta das colunas */
    .stColumn > div {
        border: 2px solid #ddd;
        padding: 20px;
        border-radius: 10px;
        background-color: #f9f9f9;
        box-shadow: 2px 2px 5px rgba(0,0,0,0.1);
    }
    </style>
""", unsafe_allow_html=True)

# Title
st.title("Agentes de IA para capacitar médicos")

# Divida a tela em duas colunas
col1, col2 = st.columns(2)

# Coluna 1: Entrada de informações
with col1:
    st.header("Entrada de Informações")# Text Inputs
    gender = st.selectbox('Selecione o gênero', ('Macho', 'Fêmea', 'Outro'))
    age = st.number_input('Entre a idade', min_value=0, max_value=120, value=25)
    symptoms = st.text_area('Entre com os sintomas', 'exemplo: febre, tosse, dor de cabeça')
    medical_history = st.text_area('Entre o histórico médico', 'exemplo: diabete, hipertensão')

# Initialize Tools
search_tool = SerperDevTool()
scrape_tool = ScrapeWebsiteTool()

llm = ChatOpenAI(
    model="gpt-3.5-turbo-16k",
    temperature=0.1,
    max_tokens=8000
)

# Define Agents
diagnostician = Agent(
    role="Diagnosticador Médico",
    goal="Analise os sintomas do paciente e o histórico médico para fornecer um diagnóstico preliminar.",
    backstory="Este agente é especializado no diagnóstico de condições médicas com base nos sintomas relatados pelo paciente e no histórico médico. Ele usa algoritmos avançados e conhecimento médico para identificar possíveis problemas de saúde.",
    verbose=True,
    allow_delegation=False,
    tools=[search_tool, scrape_tool],
    llm=llm
)

treatment_advisor = Agent(
    role="Conselheiro de Tratamento",
    goal="Recomendar planos de tratamento apropriados com base no diagnóstico fornecido pelo médico diagnosticador  em português.",
    backstory="Este agente é especializado na criação de planos de tratamento adaptados às necessidades individuais do paciente. Ele considera o diagnóstico, o histórico do paciente e as melhores práticas atuais da medicina para recomendar tratamentos eficazes.",
    verbose=True,
    allow_delegation=False,
    tools=[search_tool, scrape_tool],
    llm=llm
)

# Define Tasks
diagnose_task = Task(
    description=(
        "1. Analise os sintomas do paciente ({symptoms}) e o histórico médico ({medical_history}).\n"
        "2. Forneça um diagnóstico preliminar com possíveis condições com base nas informações fornecidas.\n"
        "3. Limite o diagnóstico às condições mais prováveis."
    ),
    expected_output="Um diagnóstico preliminar com uma lista de possíveis condições em português.",
    agent=diagnostician
)

treatment_task = Task(
    description=(
        "1. Com base no diagnóstico, recomende planos de tratamento apropriados passo a passo.\n"
        "2. Considere o histórico médico do paciente ({medical_history}) e sintomas atuais ({symptoms}).\n"
        "3. Forneça recomendações detalhadas de tratamento, incluindo medicamentos, mudanças no estilo de vida e cuidados de acompanhamento."
    ),
    expected_output="Um plano de tratamento abrangente e adaptado às necessidades do paciente em português.",
    agent=treatment_advisor
)

# Create Crew
crew = Crew(
    agents=[diagnostician, treatment_advisor],
    tasks=[diagnose_task, treatment_task],
    verbose=False
)

# Coluna 2: Saída de respostas dos agentes
with col2:
    st.header("Respostas dos Agentes")
    
    # Execution
    if st.button("Obtenha diagnóstico e plano de tratamento"):
        with st.spinner('Gerando recomendações...'):
            result = crew.kickoff(inputs={"symptoms": symptoms, "medical_history": medical_history})

            result = translate_text(str(result), source_lang='en', target_lang='pt')
            st.write(result)

            docx_file = generate_docx(result)
            download_link = get_download_link(docx_file, "Diagnóstico_e_Plano_de_Tratamento.docx")
            st.markdown(download_link, unsafe_allow_html=True)